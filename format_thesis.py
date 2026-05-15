#!/usr/bin/env python3
"""
格式化毕业论文 无图版本.docx
要求：
1. 中文宋体，西文/数字 Times New Roman；全文两侧对齐；数字后字母前空一格
2. 表格题注黑体五号，编号与名称之间两格；表内小五号，左对齐
3. 全文 CNT → CNTs
4. 图片题注五号黑体；修正半括号为完整括号
5. 3.1/3.2节 "纳米纤维素分散体" → "CNF/CNC分散液"
6. 3.3节 xxxx薄膜→CNTs/CNF/CNC薄膜；xxxx浆料→CNTs/CNF/CNC分散液；
   所有省略号占位符也替换；含图注和标题
"""
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = '/Users/zyt/claude_code/leftyeye/无图版本.docx'
OUTPUT = '/Users/zyt/claude_code/leftyeye/格式化版本.docx'

doc = Document(INPUT)

# ─────────────────────────────────────────────────────
# 工具函数
# ─────────────────────────────────────────────────────

def para_text(p):
    return ''.join(r.text for r in p.runs)

def set_para_text(p, new_text):
    """把新文本塞进第一个 run，清空其余 run（保留第一个 run 的字符格式）"""
    if not p.runs:
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ''

def apply_font(run, cn='宋体', en='Times New Roman'):
    """同时设置中文字体（eastAsia）和西文字体（ascii/hAnsi）"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:cs'), en)

def apply_hei(run, en='Times New Roman'):
    """黑体"""
    apply_font(run, cn='SimHei', en=en)

def sub(p, old, new):
    t = para_text(p)
    if old in t:
        set_para_text(p, t.replace(old, new))
        return True
    return False

def resub(p, pattern, repl, flags=0):
    t = para_text(p)
    n = re.sub(pattern, repl, t, flags=flags)
    if n != t:
        set_para_text(p, n)
        return True
    return False

# ─────────────────────────────────────────────────────
# 章节追踪
# ─────────────────────────────────────────────────────

SECTION_MAP = {
    '溶剂体系对纳米纤维素薄膜表面形貌的影响': '3.1',
    '醇醚对纳米纤维素体系结构与性能的影响':    '3.1',
    'TPM含量对纳米纤维素薄膜表面形貌的影响':   '3.2',
    'TPM含量对纳米纤维素薄膜微孔结构的调控':   '3.2',
    '碳纳米管/纳米纤维素复合薄膜的结构与性能研究': '3.3',
    '纳米纤维素基导电薄膜的结构与性能分析':    '3.3',
    '结论':   '4',
    '参考文献': 'refs',
    '致谢':   'ack',
}

cur_sec = 'front'

def update_sec(p):
    global cur_sec
    style = p.style.name if p.style else ''
    if 'Heading 1' in style or 'Heading 2' in style:
        t = para_text(p).strip()
        if t in SECTION_MAP:
            cur_sec = SECTION_MAP[t]
            print(f'  → 进入章节 {cur_sec}')

# ─────────────────────────────────────────────────────
# 各项要求的处理函数
# ─────────────────────────────────────────────────────

# ── 要求3：CNT → CNTs（独立出现时）──────────────────
# 匹配 CNT 不被字母/s/斜杠包围（避免改 CNTs、CNT/CNF 中的 /）
CNT_RE = r'(?<![A-Za-z/])CNT(?![sA-Za-z/])'

def fix_cnt(p):
    resub(p, CNT_RE, 'CNTs')
    sub(p, 'CNTS', 'CNTs')   # 修正大写拼写错误

# ── 要求5：3.1/3.2 节替换 ────────────────────────────
def fix_31_32(p):
    sub(p, '纳米纤维素分散体', 'CNF/CNC分散液')

# ── 要求6：3.3 节替换 ────────────────────────────────
# 省略号/点号占位符（中文句号、英文句点、Unicode省略号）
DOT = r'[…\.。]{2,}'

def fix_33(p):
    t = para_text(p)
    n = t
    # 省略号 + 膜/薄膜/浆料
    n = re.sub(DOT + r'膜',   'CNTs/CNF/CNC薄膜',   n)
    n = re.sub(DOT + r'薄膜', 'CNTs/CNF/CNC薄膜',   n)
    n = re.sub(DOT + r'浆料', 'CNTs/CNF/CNC分散液', n)
    # 剩余独立省略号占位符
    n = re.sub(DOT, 'CNTs/CNF/CNC薄膜', n)
    # 具名替换
    n = re.sub(r'CNTs?/CNF复合导电薄膜', 'CNTs/CNF/CNC薄膜',   n)
    n = re.sub(r'CNTs?/CNF复合薄膜',     'CNTs/CNF/CNC薄膜',   n)
    n = n.replace('CNT/CNF复合导电浆料', 'CNTs/CNF/CNC分散液')
    n = n.replace('复合导电薄膜',         'CNTs/CNF/CNC薄膜')
    n = n.replace('纤维素基导电薄膜',     'CNTs/CNF/CNC薄膜')
    n = n.replace('导电浆料',             'CNTs/CNF/CNC分散液')
    n = n.replace('CNF复合浆料',          'CNTs/CNF/CNC分散液')
    # 3.3中 "导电薄膜" 单独出现时也替换（图注、标题除外在下方专门处理）
    n = n.replace('导电薄膜',             'CNTs/CNF/CNC薄膜')
    if n != t:
        set_para_text(p, n)

# ── 要求4：修正图注半括号 ─────────────────────────────
# 半宽右括号缺左括号： a) → (a)
HALF_HW = re.compile(r'(?<![a-zA-Z（(])([a-z])\)')
# 全宽右括号缺左括号： a） → （a）
HALF_FW = re.compile(r'(?<![a-zA-Z（(])([a-z])）')

def fix_brackets(p):
    t = para_text(p)
    n = HALF_HW.sub(r'(\1)', t)
    n = HALF_FW.sub(r'（\1）', n)
    if n != t:
        set_para_text(p, n)

# ── 要求2：表题注格式（编号与名称之间两格）─────────────
TABLE_CAP = re.compile(r'^表\s*(\S+)\s+(.+)')

def fix_table_cap(p):
    t = para_text(p).strip()
    m = TABLE_CAP.match(t)
    if m:
        fixed = f'表{m.group(1)}  {m.group(2).strip()}'
        if fixed != t:
            set_para_text(p, fixed)

# ── 要求1：数字后字母前加空格 ────────────────────────────
# 只在小写字母前（单位缩写），避免影响英文缩写词内部
NUM_LETTER = re.compile(r'(\d)([a-z])')

def fix_num_letter(p):
    resub(p, NUM_LETTER, r'\1 \2')

# ─────────────────────────────────────────────────────
# 主处理循环
# ─────────────────────────────────────────────────────

print('=== 开始处理段落 ===')
for p in doc.paragraphs:
    update_sec(p)
    style  = p.style.name if p.style else ''
    raw    = para_text(p).strip()

    # ── Caption 段落（图题 / 表题）──────────────────────
    if style == 'Caption':
        if raw.startswith('图'):
            # 要求6：3.3节图注替换
            if cur_sec == '3.3':
                fix_33(p)
            # 要求5：3.1/3.2节图注
            elif cur_sec in ('3.1', '3.2'):
                fix_31_32(p)
            # 要求3：CNT → CNTs
            fix_cnt(p)
            # 要求4：修正半括号
            fix_brackets(p)
            # 要求4：五号黑体
            for r in p.runs:
                apply_hei(r)
                r.font.size = Pt(10.5)  # 五号

        elif raw.startswith('表'):
            # 要求2：表题注格式
            fix_table_cap(p)
            # 要求3：CNT → CNTs（表题注中也替换）
            fix_cnt(p)
            # 要求2：黑体五号
            for r in p.runs:
                apply_hei(r)
                r.font.size = Pt(10.5)  # 五号
        continue  # Caption 处理完毕，跳过后续

    # ── 非 Caption 段落 ──────────────────────────────────

    # 要求1：字体（所有正文段落）
    for r in p.runs:
        apply_font(r)

    # 要求1：对齐（正文段落两端对齐）
    if style in ('Normal', 'List Paragraph', 'Body Text', 'Bibliography'):
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 要求3：CNT → CNTs
    fix_cnt(p)

    # 要求1：数字后字母前空格（只对正文）
    if style in ('Normal', 'List Paragraph'):
        fix_num_letter(p)

    # 要求5：3.1/3.2 节
    if cur_sec in ('3.1', '3.2'):
        fix_31_32(p)

    # 要求6：3.3 节（含 Heading 3 标题和正文）
    if cur_sec == '3.3':
        fix_33(p)

# ─────────────────────────────────────────────────────
# 表格内容格式（要求2）
# ─────────────────────────────────────────────────────

print('=== 处理表格 ===')
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for cp in cell.paragraphs:
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
                for r in cp.runs:
                    apply_font(r)           # 字体
                    r.font.size = Pt(9)     # 小五号

# ─────────────────────────────────────────────────────
# 保存
# ─────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f'\n✓ 已保存 → {OUTPUT}')
